"""
Claude Document Agent for processing, creating and editing PDF, Excel, Word documents and images.
Uses Claude Agent SDK with AWS Bedrock backend.

Supports:
- Document analysis (PDF, Excel, Word, images)
- Document creation (PDF from HTML, Excel, Word)
- Presentation creation (HTML -> PDF)
- Streaming document generation for real-time preview

Requires: pip install claude-agent-sdk openpyxl python-docx weasyprint
"""
from __future__ import annotations

import base64
import json
import logging
import os
import tempfile
import shutil
import asyncio
from typing import Optional, AsyncIterator, Tuple
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)


class ClaudeDocumentAgent:
    """
    Agent for processing and editing documents using Claude Agent SDK.
    Uses AWS Bedrock backend for Claude models.
    """

    def __init__(self):
        """
        Initialize Claude Document Agent with AWS Bedrock backend.
        Uses AWS_BEARER_TOKEN_BEDROCK environment variable for authentication.
        """
        self.aws_bearer_token = os.getenv("AWS_BEARER_TOKEN_BEDROCK")
        self.aws_region = os.getenv("AWS_REGION", "us-east-1")

        # Set environment variables for Claude Agent SDK to use Bedrock
        if self.aws_bearer_token:
            os.environ["CLAUDE_CODE_USE_BEDROCK"] = "1"
            os.environ["AWS_REGION"] = self.aws_region
            logger.info(f"Claude Agent SDK configured with AWS Bedrock (region: {self.aws_region})")
            self._initialized = True
        else:
            logger.warning("AWS_BEARER_TOKEN_BEDROCK not set - Claude document processing unavailable")
            self._initialized = False

        # Workspace directory for file operations
        self.workspace_dir = os.getenv("CLAUDE_WORKSPACE_DIR", "./claude_workspace")
        os.makedirs(self.workspace_dir, exist_ok=True)

    async def process_document(
        self,
        file_bytes: bytes,
        file_name: str,
        mime_type: str,
        task: str,
        language: str = "ru",
    ) -> str:
        """
        Process a document using AWS Bedrock Converse API directly.

        Args:
            file_bytes: Document file bytes
            file_name: Original filename
            mime_type: MIME type of the document
            task: What to do with the document
            language: Response language (ru/en)

        Returns:
            Processing result as text
        """
        if not self._initialized:
            return "Ошибка: AWS Bedrock не настроен. Установите AWS_BEARER_TOKEN_BEDROCK."

        # Use boto3 directly for reliable document processing
        return await self._process_with_boto3(file_bytes, file_name, mime_type, task, language)

    async def _process_with_boto3(
        self,
        file_bytes: bytes,
        file_name: str,
        mime_type: str,
        task: str,
        language: str = "ru",
    ) -> str:
        """
        Process document using boto3 direct API with AWS Bedrock.
        """
        logger.info(f"📄 [BOTO3] Processing document: {file_name} ({mime_type})")
        logger.info(f"📄 [BOTO3] File size: {len(file_bytes)} bytes")
        logger.info(f"📄 [BOTO3] Task: {task[:100]}...")

        try:
            import boto3
            import re

            # Create Bedrock client
            logger.info(f"📄 [BOTO3] Creating Bedrock client for region: {self.aws_region}")
            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            # Determine document type
            doc_type = self._get_document_type(mime_type, file_name)
            system_prompt = self._build_system_prompt(doc_type, language)

            # Build content based on file type
            is_image = mime_type.startswith("image/")

            # Use Claude Haiku 4.5 for all content (supports both images and documents)
            model_id = "us.anthropic.claude-haiku-4-5-20251001-v1:0"

            if is_image:
                image_format = mime_type.split("/")[-1]
                if image_format == "jpg":
                    image_format = "jpeg"

                content = [
                    {
                        "image": {
                            "format": image_format,
                            "source": {
                                "bytes": file_bytes
                            }
                        }
                    },
                    {
                        "text": f"Задача: {task}\n\nФайл: {file_name}"
                    }
                ]
            else:

                doc_format = self._get_doc_format(mime_type, file_name)

                # Sanitize filename
                sanitized_name = re.sub(r'[^a-zA-Z0-9\s\-\(\)\[\]]', '_', file_name)
                sanitized_name = re.sub(r'\s+', ' ', sanitized_name)
                sanitized_name = sanitized_name.strip()[:100]

                content = [
                    {
                        "document": {
                            "format": doc_format,
                            "name": sanitized_name,
                            "source": {
                                "bytes": file_bytes
                            }
                        }
                    },
                    {
                        "text": f"Задача: {task}"
                    }
                ]

            logger.info(f"Calling Bedrock Converse API with model: {model_id}")

            # Call Bedrock Converse API
            import asyncio
            loop = asyncio.get_event_loop()

            response = await loop.run_in_executor(
                None,
                lambda: client.converse(
                    modelId=model_id,
                    messages=[
                        {
                            "role": "user",
                            "content": content
                        }
                    ],
                    system=[
                        {"text": system_prompt}
                    ],
                    inferenceConfig={
                        "maxTokens": 4096,
                        "temperature": 0
                    }
                )
            )

            # Extract text from response
            result_text = ""
            output = response.get("output", {})
            message = output.get("message", {})
            for block in message.get("content", []):
                if "text" in block:
                    result_text += block.get("text", "")

            logger.info(f"Boto3 response received, length: {len(result_text)}")
            return result_text or "Документ обработан, но результат пуст."

        except Exception as e:
            logger.error(f"Boto3 Bedrock API error: {e}", exc_info=True)
            return f"Ошибка Bedrock API: {str(e)}"

    async def edit_document(
        self,
        file_bytes: bytes,
        file_name: str,
        mime_type: str,
        edit_instructions: str,
        language: str = "ru",
    ) -> tuple[bytes, str]:
        """
        Edit a document using Claude Agent SDK.

        Args:
            file_bytes: Original document bytes
            file_name: Original filename
            mime_type: MIME type of the document
            edit_instructions: Instructions for editing the document
            language: Response language (ru/en)

        Returns:
            Tuple of (edited_file_bytes, summary_of_changes)
        """
        if not self._initialized:
            return file_bytes, "Ошибка: AWS Bedrock не настроен."

        try:
            from claude_agent_sdk import query, ClaudeAgentOptions
            import re

            # Sanitize filename
            sanitized_name = re.sub(r'[^a-zA-Z0-9\s\-\(\)\[\]._]', '_', file_name)
            sanitized_name = re.sub(r'\s+', ' ', sanitized_name)
            sanitized_name = sanitized_name.strip()[:100]

            # Create temp directory for this edit session
            session_dir = tempfile.mkdtemp(prefix="claude_edit_", dir=self.workspace_dir)
            file_path = os.path.join(session_dir, sanitized_name)

            # Save original file
            with open(file_path, "wb") as f:
                f.write(file_bytes)

            logger.info(f"Saved file for editing: {file_path}")

            # Build prompt for editing
            lang_instruction = (
                "Отвечай на русском языке." if language == "ru" else "Respond in English."
            )

            prompt = f"""{lang_instruction}

Файл для редактирования: {file_path}
Инструкции: {edit_instructions}

1. Прочитай файл с помощью Read
2. Внеси необходимые изменения с помощью Edit или Write
3. Опиши что было изменено"""

            # Configure options with write permissions
            options = ClaudeAgentOptions(
                system_prompt=f"""Ты - ассистент для редактирования документов.
{lang_instruction}
Ты умеешь читать и редактировать файлы.""",
                max_turns=5,
                allowed_tools=["Read", "Write", "Edit", "Glob"],
            )

            logger.info(f"Calling Claude Agent SDK to edit: {sanitized_name}")

            # Run Claude Agent for editing
            summary = ""
            async for message in query(prompt=prompt, options=options):
                if hasattr(message, 'content'):
                    if isinstance(message.content, str):
                        summary += message.content
                    elif isinstance(message.content, list):
                        for block in message.content:
                            if hasattr(block, 'text'):
                                summary += block.text
                elif hasattr(message, 'text'):
                    summary += message.text
                elif isinstance(message, str):
                    summary += message

            # Read edited file
            edited_bytes = file_bytes  # Default to original if edit failed
            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    edited_bytes = f.read()

            # Cleanup session directory
            try:
                shutil.rmtree(session_dir)
                logger.info(f"Cleaned up edit session: {session_dir}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to cleanup edit session: {cleanup_error}")

            return edited_bytes, summary or "Файл обработан."

        except ImportError:
            logger.error("claude-agent-sdk not installed for editing")
            return file_bytes, "Ошибка: claude-agent-sdk не установлен для редактирования файлов."
        except Exception as e:
            logger.error(f"Document edit error: {e}", exc_info=True)
            return file_bytes, f"Ошибка при редактировании: {str(e)}"

    def _get_document_type(self, mime_type: str, filename: str) -> str:
        """Determine document type from MIME type and filename."""
        filename_lower = filename.lower()

        if "pdf" in mime_type or filename_lower.endswith(".pdf"):
            return "pdf"
        elif "spreadsheet" in mime_type or filename_lower.endswith((".xlsx", ".xls", ".csv")):
            return "excel"
        elif "word" in mime_type or "document" in mime_type or filename_lower.endswith((".docx", ".doc")):
            return "word"
        elif "text" in mime_type or filename_lower.endswith(".txt"):
            return "text"
        elif mime_type.startswith("image/"):
            return "image"
        else:
            return "unknown"

    def _get_doc_format(self, mime_type: str, filename: str) -> str:
        """Get document format for Bedrock Converse API."""
        filename_lower = filename.lower()

        if "pdf" in mime_type or filename_lower.endswith(".pdf"):
            return "pdf"
        elif filename_lower.endswith(".docx"):
            return "docx"
        elif filename_lower.endswith(".doc"):
            return "doc"
        elif filename_lower.endswith(".xlsx"):
            return "xlsx"
        elif filename_lower.endswith(".xls"):
            return "xls"
        elif filename_lower.endswith(".csv"):
            return "csv"
        elif filename_lower.endswith(".txt"):
            return "txt"
        elif filename_lower.endswith(".html"):
            return "html"
        elif filename_lower.endswith(".md"):
            return "md"
        else:
            return "txt"

    def _build_system_prompt(self, doc_type: str, language: str) -> str:
        """Build system prompt based on document type and language."""
        lang_instruction = (
            "Отвечай на русском языке." if language == "ru" else "Respond in English."
        )

        base_prompt = f"""Ты - специализированный ассистент для анализа и редактирования документов.
Используешь Claude Agent SDK для работы с файлами.

{lang_instruction}

"""

        prompts = {
            "pdf": base_prompt + """Ты специализируешься на работе с PDF документами.
Умеешь:
- Извлекать текст и таблицы
- Анализировать структуру документа
- Находить ключевую информацию
- Суммировать содержание
- Отвечать на вопросы о содержимом

Предоставляй точные и структурированные ответы.""",

            "excel": base_prompt + """Ты специализируешься на работе с таблицами Excel и CSV.
Умеешь:
- Анализировать данные в таблицах
- Вычислять статистику
- Находить паттерны и тренды
- Создавать сводки данных
- Редактировать данные в ячейках
- Добавлять новые строки/столбцы

Предоставляй числовые результаты с точностью.""",

            "word": base_prompt + """Ты специализируешься на работе с Word документами.
Умеешь:
- Анализировать текстовые документы
- Извлекать ключевую информацию
- Суммировать содержание
- Редактировать текст
- Добавлять и удалять разделы

Сохраняй форматирование и структуру при необходимости.""",

            "image": base_prompt + """Ты специализируешься на анализе изображений.
Умеешь:
- Описывать содержимое изображений
- Распознавать текст на изображениях (OCR)
- Анализировать диаграммы и графики
- Извлекать данные из скриншотов
- Отвечать на вопросы об изображении

Предоставляй детальные и точные описания.""",
        }

        return prompts.get(doc_type, base_prompt + """Ты - универсальный ассистент для работы с документами.
Анализируй и редактируй содержимое согласно запросу пользователя.""")

    async def analyze_image(
        self, file_bytes: bytes, file_name: str, mime_type: str, task: str = None
    ) -> str:
        """Analyze an image using Claude's vision capabilities."""
        if not task:
            task = "Опиши что изображено на этой картинке. Укажи все важные детали."
        return await self.process_document(
            file_bytes=file_bytes,
            file_name=file_name,
            mime_type=mime_type,
            task=task,
        )

    async def extract_text_from_image(
        self, file_bytes: bytes, file_name: str, mime_type: str
    ) -> str:
        """Extract text from an image (OCR)."""
        return await self.process_document(
            file_bytes=file_bytes,
            file_name=file_name,
            mime_type=mime_type,
            task="Извлеки весь текст с этого изображения. Сохрани форматирование где возможно.",
        )

    async def summarize_document(
        self, file_bytes: bytes, file_name: str, mime_type: str
    ) -> str:
        """Create a summary of the document."""
        return await self.process_document(
            file_bytes=file_bytes,
            file_name=file_name,
            mime_type=mime_type,
            task="Создай подробное резюме этого документа.",
        )

    # ─────────────────────────────────────────────────────────────────────
    # Document Creation Methods
    # ─────────────────────────────────────────────────────────────────────

    async def create_presentation_html(
        self,
        topic: str,
        requirements: str = "",
        language: str = "ru",
    ) -> AsyncIterator[Tuple[str, str]]:
        """
        Create a presentation as HTML with streaming.
        Yields (content_type, content) tuples:
        - ("html_chunk", html_content) for streaming HTML
        - ("complete", full_html) when done

        Args:
            topic: Presentation topic
            requirements: Additional requirements (slides count, style, etc.)
            language: Response language

        Yields:
            Tuples of (content_type, content)
        """
        logger.info(f"📊 [PRESENTATION] Creating presentation for: {topic}")

        lang_instruction = "на русском языке" if language == "ru" else "in English"

        prompt = f"""Создай HTML презентацию {lang_instruction} на тему: "{topic}"

{requirements}

ВАЖНО: Создай полноценную HTML презентацию со следующими требованиями:
1. Используй современный CSS с градиентами и анимациями
2. Каждый слайд в отдельном <section class="slide">
3. Добавь навигацию между слайдами
4. Используй красивые шрифты (Google Fonts)
5. Добавь иконки где уместно (Font Awesome или эмодзи)
6. Размер слайда: 1920x1080 (16:9)
7. Минимум 5-7 слайдов

Верни ТОЛЬКО HTML код, без пояснений. Начни с <!DOCTYPE html>"""

        try:
            import boto3

            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            model_id = "us.anthropic.claude-sonnet-4-20250514-v1:0"

            # Use streaming for real-time preview
            response = client.converse_stream(
                modelId=model_id,
                messages=[{"role": "user", "content": [{"text": prompt}]}],
                system=[{"text": "Ты - эксперт по созданию презентаций. Создавай красивые, современные HTML презентации."}],
                inferenceConfig={"maxTokens": 8192, "temperature": 0.7}
            )

            full_html = ""
            for event in response.get("stream", []):
                if "contentBlockDelta" in event:
                    delta = event["contentBlockDelta"].get("delta", {})
                    if "text" in delta:
                        chunk = delta["text"]
                        full_html += chunk
                        yield ("html_chunk", chunk)

            yield ("complete", full_html)
            logger.info(f"✅ [PRESENTATION] Created HTML presentation: {len(full_html)} chars")

        except Exception as e:
            logger.error(f"❌ [PRESENTATION] Error: {e}", exc_info=True)
            yield ("error", f"Ошибка создания презентации: {str(e)}")

    async def html_to_pdf(self, html_content: str, output_filename: str = "presentation.pdf") -> bytes:
        """
        Convert HTML to PDF using WeasyPrint.

        Args:
            html_content: HTML content to convert
            output_filename: Output filename (for logging)

        Returns:
            PDF file bytes
        """
        logger.info(f"📄 [HTML->PDF] Converting to PDF: {output_filename}")

        try:
            from weasyprint import HTML, CSS

            # Create PDF in memory
            loop = asyncio.get_event_loop()
            pdf_bytes = await loop.run_in_executor(
                None,
                lambda: HTML(string=html_content).write_pdf()
            )

            logger.info(f"✅ [HTML->PDF] Created PDF: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except ImportError:
            logger.error("WeasyPrint not installed. Install with: pip install weasyprint")
            raise ImportError("WeasyPrint required for PDF conversion. Install: pip install weasyprint")
        except Exception as e:
            logger.error(f"❌ [HTML->PDF] Error: {e}", exc_info=True)
            raise

    async def create_excel_document(
        self,
        task: str,
        data: Optional[dict] = None,
        language: str = "ru",
    ) -> AsyncIterator[Tuple[str, any]]:
        """
        Create or modify Excel document with streaming status updates.

        Args:
            task: What to create (e.g., "Создай таблицу расходов за месяц")
            data: Optional existing data to include
            language: Response language

        Yields:
            Tuples of (status, content):
            - ("status", "Generating structure...")
            - ("data", {"sheets": [...], "rows": [...]})
            - ("complete", excel_bytes)
        """
        logger.info(f"📊 [EXCEL] Creating Excel document: {task}")

        yield ("status", "Анализирую задачу..." if language == "ru" else "Analyzing task...")

        lang_instruction = "на русском языке" if language == "ru" else "in English"

        prompt = f"""Задача: "{task}"

Создай структуру Excel таблицы {lang_instruction}.

КРИТИЧНО:
- Ответь ТОЛЬКО валидным JSON объектом
- Никакого текста до или после JSON
- МАКСИМУМ 15-20 строк данных (rows), не больше!
- Ответ должен быть компактным

Формат:
{{"sheets":[{{"name":"Лист1","columns":["Столбец1","Столбец2"],"rows":[["значение1","значение2"]],"column_widths":[15,20],"header_style":{{"bold":true,"bg_color":"#4472C4","font_color":"#FFFFFF"}}}}],"title":"Название"}}"""

        try:
            import boto3

            yield ("status", "Генерирую структуру таблицы..." if language == "ru" else "Generating table structure...")

            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            model_id = "us.anthropic.claude-sonnet-4-20250514-v1:0"

            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: client.converse(
                    modelId=model_id,
                    messages=[{"role": "user", "content": [{"text": prompt}]}],
                    system=[{"text": "Ты генератор компактного JSON для Excel. Отвечай ТОЛЬКО валидным JSON. Максимум 15-20 строк данных. Без пояснений."}],
                    inferenceConfig={"maxTokens": 8192, "temperature": 0.2}
                )
            )

            # Extract JSON from response
            result_text = ""
            output = response.get("output", {})
            message = output.get("message", {})
            for block in message.get("content", []):
                if "text" in block:
                    result_text += block.get("text", "")

            # Parse JSON
            import re
            logger.info(f"📊 [EXCEL] Raw response length: {len(result_text)} chars")
            logger.info(f"📊 [EXCEL] Raw response start: {result_text[:300]}...")

            # Clean up response - remove markdown code blocks if present
            cleaned = re.sub(r'```json\s*', '', result_text)
            cleaned = re.sub(r'```\s*', '', cleaned)
            cleaned = cleaned.strip()

            # Try to find complete JSON
            json_match = re.search(r'\{[\s\S]*\}', cleaned)

            excel_structure = None

            if json_match:
                try:
                    excel_structure = json.loads(json_match.group())
                    logger.info(f"✅ [EXCEL] Parsed complete JSON")
                except json.JSONDecodeError as e:
                    logger.warning(f"⚠️ [EXCEL] JSON parse error: {e}")
                    # Try to fix common issues
                    json_str = json_match.group()
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    try:
                        excel_structure = json.loads(json_str)
                    except:
                        pass

            # If still no valid JSON, try to repair truncated JSON
            if not excel_structure:
                logger.warning(f"⚠️ [EXCEL] Attempting to repair truncated JSON...")
                # Find the start of JSON
                start_idx = cleaned.find('{')
                if start_idx >= 0:
                    json_str = cleaned[start_idx:]
                    # Count brackets to find valid cutoff point
                    depth = 0
                    last_valid_idx = 0
                    in_string = False
                    escape_next = False

                    for i, char in enumerate(json_str):
                        if escape_next:
                            escape_next = False
                            continue
                        if char == '\\':
                            escape_next = True
                            continue
                        if char == '"' and not escape_next:
                            in_string = not in_string
                            continue
                        if in_string:
                            continue
                        if char == '{' or char == '[':
                            depth += 1
                        elif char == '}' or char == ']':
                            depth -= 1
                            if depth == 0:
                                last_valid_idx = i + 1
                                break

                    if last_valid_idx > 0:
                        try:
                            excel_structure = json.loads(json_str[:last_valid_idx])
                            logger.info(f"✅ [EXCEL] Repaired truncated JSON")
                        except:
                            pass

                    # Last resort: try to close arrays/objects manually
                    if not excel_structure:
                        json_str = cleaned[start_idx:]
                        # Remove incomplete last row
                        last_bracket = json_str.rfind(']')
                        if last_bracket > 0:
                            json_str = json_str[:last_bracket+1]
                            # Close remaining structures
                            open_braces = json_str.count('{') - json_str.count('}')
                            open_brackets = json_str.count('[') - json_str.count(']')
                            json_str += ']' * open_brackets + '}' * open_braces
                            try:
                                excel_structure = json.loads(json_str)
                                logger.info(f"✅ [EXCEL] Fixed JSON with manual closing")
                            except Exception as fix_err:
                                logger.error(f"❌ [EXCEL] Could not fix JSON: {fix_err}")

            if not excel_structure:
                logger.error(f"❌ [EXCEL] No valid JSON found in: {result_text[:500]}")
                raise ValueError("No valid JSON found in response")
            yield ("data", excel_structure)

            yield ("status", "Создаю Excel файл..." if language == "ru" else "Creating Excel file...")

            # Create Excel file
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import io

            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet

            for sheet_data in excel_structure.get("sheets", []):
                ws = wb.create_sheet(title=sheet_data.get("name", "Sheet1")[:31])

                # Header style
                header_style = sheet_data.get("header_style", {})
                header_font = Font(
                    bold=header_style.get("bold", True),
                    color=header_style.get("font_color", "FFFFFF").replace("#", "")
                )
                header_fill = PatternFill(
                    start_color=header_style.get("bg_color", "4472C4").replace("#", ""),
                    end_color=header_style.get("bg_color", "4472C4").replace("#", ""),
                    fill_type="solid"
                )

                # Write headers
                columns = sheet_data.get("columns", [])
                for col_idx, col_name in enumerate(columns, 1):
                    cell = ws.cell(row=1, column=col_idx, value=col_name)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")

                # Write data rows
                for row_idx, row_data in enumerate(sheet_data.get("rows", []), 2):
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)

                # Set column widths
                for col_idx, width in enumerate(sheet_data.get("column_widths", []), 1):
                    ws.column_dimensions[get_column_letter(col_idx)].width = width

            # Save to bytes
            excel_buffer = io.BytesIO()
            wb.save(excel_buffer)
            excel_bytes = excel_buffer.getvalue()

            yield ("complete", excel_bytes)
            logger.info(f"✅ [EXCEL] Created Excel: {len(excel_bytes)} bytes")

        except ImportError as e:
            logger.error(f"openpyxl not installed: {e}")
            yield ("error", "openpyxl required. Install: pip install openpyxl")
        except Exception as e:
            logger.error(f"❌ [EXCEL] Error: {e}", exc_info=True)
            yield ("error", f"Ошибка создания Excel: {str(e)}")

    async def create_word_document(
        self,
        task: str,
        language: str = "ru",
    ) -> AsyncIterator[Tuple[str, any]]:
        """
        Create Word document with streaming content generation.

        Args:
            task: What to create (e.g., "Создай договор аренды")
            language: Response language

        Yields:
            Tuples of (status, content):
            - ("text_chunk", text) for streaming text preview
            - ("status", status_message)
            - ("complete", docx_bytes)
        """
        logger.info(f"📝 [WORD] Creating Word document: {task}")

        lang_instruction = "на русском языке" if language == "ru" else "in English"

        prompt = f"""Напиши ПОЛНЫЙ ТЕКСТ документа {lang_instruction} по запросу: "{task}"

ВАЖНО:
- Пиши ТОЛЬКО сам текст документа, который будет внутри файла
- НЕ пиши инструкции по созданию документа
- НЕ пиши "вот текст документа" или подобные фразы
- Сразу начинай с содержимого документа (заголовка)

Форматирование:
- # для основного заголовка документа
- ## для разделов
- ### для подразделов
- **жирный** для выделения важного
- - для маркированных списков
- 1. для нумерованных списков

Создай полноценный, профессиональный текст документа с реальным содержанием."""

        try:
            import boto3

            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            model_id = "us.anthropic.claude-sonnet-4-20250514-v1:0"

            # Stream the content
            response = client.converse_stream(
                modelId=model_id,
                messages=[{"role": "user", "content": [{"text": prompt}]}],
                system=[{"text": "Ты - писатель документов. Твоя задача - писать ТЕКСТ СОДЕРЖИМОГО документа, а не инструкции по его созданию. Никогда не пиши 'вот документ' или 'ниже приведен текст' - сразу начинай с заголовка документа."}],
                inferenceConfig={"maxTokens": 8192, "temperature": 0.7}
            )

            full_text = ""
            for event in response.get("stream", []):
                if "contentBlockDelta" in event:
                    delta = event["contentBlockDelta"].get("delta", {})
                    if "text" in delta:
                        chunk = delta["text"]
                        full_text += chunk
                        yield ("text_chunk", chunk)

            yield ("status", "Создаю Word файл..." if language == "ru" else "Creating Word file...")

            # Create Word document
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            import re

            doc = Document()

            # Parse markdown-like content
            lines = full_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    # Main heading
                    p = doc.add_heading(line[2:], level=0)
                elif line.startswith('## '):
                    # Subheading
                    p = doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:], level=2)
                elif line.startswith('- ') or line.startswith('* '):
                    # List item
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    # Numbered list
                    p = doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
                else:
                    # Regular paragraph
                    p = doc.add_paragraph()
                    # Handle bold text
                    parts = re.split(r'(\*\*[^*]+\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()

            yield ("complete", docx_bytes)
            logger.info(f"✅ [WORD] Created Word document: {len(docx_bytes)} bytes")

        except ImportError as e:
            logger.error(f"python-docx not installed: {e}")
            yield ("error", "python-docx required. Install: pip install python-docx")
        except Exception as e:
            logger.error(f"❌ [WORD] Error: {e}", exc_info=True)
            yield ("error", f"Ошибка создания Word: {str(e)}")

    async def modify_excel_document(
        self,
        file_bytes: bytes,
        file_name: str,
        instructions: str,
        language: str = "ru",
    ) -> AsyncIterator[Tuple[str, any]]:
        """
        Modify existing Excel document.

        Args:
            file_bytes: Original Excel file bytes
            file_name: Original filename
            instructions: Modification instructions
            language: Response language

        Yields:
            Status updates and final modified file
        """
        logger.info(f"📊 [EXCEL-EDIT] Modifying: {file_name}")

        yield ("status", "Читаю файл..." if language == "ru" else "Reading file...")

        try:
            from openpyxl import load_workbook
            import io

            # Load workbook
            wb = load_workbook(io.BytesIO(file_bytes))

            # Get current structure for AI analysis
            structure = {"sheets": []}
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_info = {
                    "name": sheet_name,
                    "rows": [],
                    "max_row": ws.max_row,
                    "max_col": ws.max_column
                }
                # Get first 20 rows as sample
                for row in ws.iter_rows(min_row=1, max_row=min(20, ws.max_row), values_only=True):
                    sheet_info["rows"].append(list(row))
                structure["sheets"].append(sheet_info)

            yield ("status", "Анализирую изменения..." if language == "ru" else "Analyzing changes...")

            # Ask AI what changes to make
            import boto3

            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            prompt = f"""Текущая структура Excel файла:
{json.dumps(structure, ensure_ascii=False, indent=2)}

Инструкции для изменения: {instructions}

Верни JSON с изменениями в формате:
{{
    "changes": [
        {{"sheet": "Sheet1", "cell": "A1", "value": "новое значение"}},
        {{"sheet": "Sheet1", "action": "add_row", "row": 5, "values": ["a", "b", "c"]}},
        {{"sheet": "Sheet1", "action": "delete_row", "row": 3}}
    ],
    "summary": "Описание изменений"
}}

Верни ТОЛЬКО JSON."""

            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: client.converse(
                    modelId="us.anthropic.claude-sonnet-4-20250514-v1:0",
                    messages=[{"role": "user", "content": [{"text": prompt}]}],
                    inferenceConfig={"maxTokens": 4096, "temperature": 0}
                )
            )

            result_text = ""
            for block in response.get("output", {}).get("message", {}).get("content", []):
                if "text" in block:
                    result_text += block.get("text", "")

            import re
            json_match = re.search(r'\{[\s\S]*\}', result_text)
            if json_match:
                changes_data = json.loads(json_match.group())

                yield ("status", "Применяю изменения..." if language == "ru" else "Applying changes...")

                # Apply changes
                for change in changes_data.get("changes", []):
                    sheet_name = change.get("sheet", wb.sheetnames[0])
                    if sheet_name not in wb.sheetnames:
                        continue
                    ws = wb[sheet_name]

                    if "cell" in change:
                        ws[change["cell"]] = change.get("value")
                    elif change.get("action") == "add_row":
                        ws.insert_rows(change.get("row", ws.max_row + 1))
                        for col, val in enumerate(change.get("values", []), 1):
                            ws.cell(row=change.get("row"), column=col, value=val)
                    elif change.get("action") == "delete_row":
                        ws.delete_rows(change.get("row"))

                yield ("data", {"summary": changes_data.get("summary", "Изменения применены")})

            # Save modified workbook
            excel_buffer = io.BytesIO()
            wb.save(excel_buffer)

            yield ("complete", excel_buffer.getvalue())
            logger.info(f"✅ [EXCEL-EDIT] Modified successfully")

        except Exception as e:
            logger.error(f"❌ [EXCEL-EDIT] Error: {e}", exc_info=True)
            yield ("error", f"Ошибка изменения Excel: {str(e)}")

    async def modify_word_document(
        self,
        file_bytes: bytes,
        file_name: str,
        instructions: str,
        language: str = "ru",
    ) -> AsyncIterator[Tuple[str, any]]:
        """
        Modify existing Word document with streaming preview.
        """
        logger.info(f"📝 [WORD-EDIT] Modifying: {file_name}")

        yield ("status", "Читаю документ..." if language == "ru" else "Reading document...")

        try:
            from docx import Document
            import io

            # Load document
            doc = Document(io.BytesIO(file_bytes))

            # Extract current content
            current_text = ""
            for para in doc.paragraphs:
                current_text += para.text + "\n"

            yield ("status", "Генерирую изменения..." if language == "ru" else "Generating changes...")

            # Ask AI for modified content
            import boto3

            client = boto3.client(
                service_name="bedrock-runtime",
                region_name=self.aws_region,
            )

            prompt = f"""Текущее содержимое документа:
---
{current_text[:5000]}
---

Инструкции для изменения: {instructions}

Создай измененную версию документа. Используй форматирование:
- # для заголовков
- ## для подзаголовков
- **жирный** для важного

Верни ТОЛЬКО измененный текст документа."""

            response = client.converse_stream(
                modelId="us.anthropic.claude-sonnet-4-20250514-v1:0",
                messages=[{"role": "user", "content": [{"text": prompt}]}],
                inferenceConfig={"maxTokens": 8192, "temperature": 0.3}
            )

            modified_text = ""
            for event in response.get("stream", []):
                if "contentBlockDelta" in event:
                    delta = event["contentBlockDelta"].get("delta", {})
                    if "text" in delta:
                        chunk = delta["text"]
                        modified_text += chunk
                        yield ("text_chunk", chunk)

            yield ("status", "Сохраняю документ..." if language == "ru" else "Saving document...")

            # Create new document with modified content
            new_doc = Document()
            import re

            for line in modified_text.split('\n'):
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    new_doc.add_heading(line[2:], level=0)
                elif line.startswith('## '):
                    new_doc.add_heading(line[3:], level=1)
                elif line.startswith('- '):
                    new_doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    p = new_doc.add_paragraph()
                    parts = re.split(r'(\*\*[^*]+\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

            # Save to bytes
            docx_buffer = io.BytesIO()
            new_doc.save(docx_buffer)

            yield ("complete", docx_buffer.getvalue())
            logger.info(f"✅ [WORD-EDIT] Modified successfully")

        except Exception as e:
            logger.error(f"❌ [WORD-EDIT] Error: {e}", exc_info=True)
            yield ("error", f"Ошибка изменения Word: {str(e)}")


# Global singleton instance
_claude_agent: Optional[ClaudeDocumentAgent] = None


def get_claude_agent() -> ClaudeDocumentAgent:
    """Get or create Claude Document Agent instance."""
    global _claude_agent
    if _claude_agent is None:
        _claude_agent = ClaudeDocumentAgent()
    return _claude_agent
